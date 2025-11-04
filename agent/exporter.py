from docx import Document
from docx.shared import Pt

class Exporter:
    @staticmethod
    def to_docx(markdown_text: str, out_path: str):
        """
        Простой и стабильный экспорт: грубо конвертим заголовки (#, ##, ###) и абзацы.
        Этого достаточно для согласований. Без сторонних конвертеров.
        """
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for block in markdown_text.split("\n\n"):
            b = block.strip()
            if not b:
                continue
            # заголовки
            if b.startswith("#"):
                level = len(b) - len(b.lstrip("#"))
                text = b.lstrip("# ").strip()
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
                continue
            # обычный параграф
            doc.add_paragraph(b)

        doc.save(out_path)