"""Document parsing utilities for the agent platform."""

from __future__ import annotations

import html
import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Sequence

from .contracts import CatalogItem, DocumentInput, ParsedDocument

logger = logging.getLogger(__name__)

EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"\+?\d[\d\-() ]{7,}\d")
UPPERCASE_HEADER_RE = re.compile(r"^([A-ZА-Я][A-ZА-Я0-9 \-/]{3,})\s*(?:®|™)?$")


def mask_pii(value: str) -> str:
    """Mask common PII artefacts such as emails and phone numbers."""

    value = EMAIL_RE.sub("[redacted-email]", value)
    value = PHONE_RE.sub("[redacted-phone]", value)
    return value


def chunk_text(text: str, max_chars: int = 800) -> List[str]:
    """
    Greedy chunking by words: накапливаем слова, пока длина чанка
    НЕ ДОСТИГНЕТ или НЕ ПРЕВЫСИТ max_chars; после этого закрываем чанк.
    Пример: chunk_text("one two three four", 7) -> ["one two", "three four"]
    """
    if not text:
        return []

    words = text.split()
    chunks: List[str] = []
    cur: List[str] = []
    cur_len = 0  # длина ' '.join(cur)

    for w in words:
        if not cur:
            # начинаем новый чанк
            cur.append(w)
            cur_len = len(w)
            # если одно слово уже >= лимита — закрываем сразу
            if cur_len >= max_chars:
                chunks.append(" ".join(cur))
                cur, cur_len = [], 0
            continue

        proposed_len = cur_len + 1 + len(w)  # +1 за пробел
        if proposed_len < max_chars:
            # ниже лимита — продолжаем накапливать
            cur.append(w)
            cur_len = proposed_len
        else:
            # достигли/превысили лимит — добавляем слово и закрываем чанк
            cur.append(w)
            chunks.append(" ".join(cur))
            cur, cur_len = [], 0

    if cur:
        chunks.append(" ".join(cur))
    return chunks

def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_docx_file(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency issue
        raise RuntimeError("python-docx must be installed to parse DOCX files") from exc

    document = Document(str(path))
    paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
    return "\n".join(paragraphs)


def _read_pdf_file(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency issue
        raise RuntimeError("pypdf must be installed to parse PDF files") from exc

    reader = PdfReader(str(path))
    content: List[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        content.append(text.strip())
    return "\n".join(filter(None, content))


READERS = {
    "text": _read_text_file,
    "markdown": _read_text_file,
    "docx": _read_docx_file,
    "pdf": _read_pdf_file,
}


@dataclass
class DocumentParser:
    """Parse structured and semi-structured reference documents."""

    chunk_size: int = 800

    def parse(self, inputs: Sequence[DocumentInput]) -> List[ParsedDocument]:
        """Parse a sequence of documents into structured records."""

        parsed: List[ParsedDocument] = []
        for spec in inputs:
            path = Path(spec.path)
            if not path.exists():
                logger.warning("Document not found", extra={"path": str(path)})
                continue
            reader = READERS.get(spec.type)
            if reader is None:
                logger.warning("Unsupported document type", extra={"type": spec.type})
                continue
            raw_text = reader(path)
            sanitized = mask_pii(html.unescape(raw_text))
            chunks = chunk_text(sanitized, max_chars=self.chunk_size)
            parsed.append(
                ParsedDocument(
                    label=spec.label,
                    text=sanitized,
                    metadata={"path": str(path.resolve()), "type": spec.type},
                    chunks=chunks,
                )
            )
        return parsed

    def extract_catalog(self, document: ParsedDocument) -> Dict[str, CatalogItem]:
        """Extract catalogue items from a parsed document."""

        lines = [line.strip() for line in document.text.splitlines() if line.strip()]
        catalog: Dict[str, CatalogItem] = {}
        current: CatalogItem | None = None
        for line in lines:
            header = UPPERCASE_HEADER_RE.match(line)
            if header:
                if current is not None:
                    catalog[current.name] = current
                name = header.group(1).strip()
                current = CatalogItem(name=name.title(), tagline="")
                continue
            if current is None:
                continue
            if not current.tagline:
                current.tagline = line
                continue
            if line.startswith("-") or line.startswith("•") or line.startswith("·"):
                benefit = line.lstrip("-•· ")
                if benefit:
                    current.benefits.append(benefit)
                continue
            if any(token in line.lower() for token in ["дет", "взрос", "сем", "family", "adult", "child"]):
                current.target_audience.append(line)
            current.notes = (current.notes + " " + line).strip()
        if current is not None:
            catalog[current.name] = current
        return catalog

    def extract_catalog_from_many(self, documents: Iterable[ParsedDocument]) -> Dict[str, CatalogItem]:
        """Aggregate catalogue items from multiple documents."""

        aggregate: Dict[str, CatalogItem] = {}
        for doc in documents:
            items = self.extract_catalog(doc)
            for name, item in items.items():
                aggregate[name] = item
        return aggregate


def convert_docx_directory(directory: Path) -> List[Path]:
    """Convert all DOCX files in a directory to ``.txt`` files."""

    created: List[Path] = []
    for path in directory.glob("*.docx"):
        text = _read_docx_file(path)
        output = path.with_suffix(path.suffix + ".txt")
        output.write_text(text, encoding="utf-8")
        created.append(output)
    return created
